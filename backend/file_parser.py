"""
file_parser.py
--------------
Extracts plain text from an uploaded essay file (.pdf, .docx, or .txt),
so the user can upload a document instead of pasting text.

Design goals:
- Never crash the server on a bad file. Corrupted files, password-protected
  PDFs, empty documents, and docx files containing images all need to fail
  *gracefully* with a clear message — not with a raw stack trace.
- Images inside a .docx are simply skipped when extracting text (python-docx
  only reads paragraph text runs, never image data, so this is naturally
  safe) — but we count and report how many were skipped, so the user knows
  something was left out rather than silently losing content.
- Formatting: paragraph breaks are preserved (joined with blank lines,
  matching how features.py's paragraph_count already expects text to be
  structured), but this deliberately does NOT try to preserve rich
  formatting like bold/italic/tables — the model only reads plain text
  anyway, so keeping that would add complexity with no analytical benefit.
"""

import io
from pypdf import PdfReader
from pypdf.errors import PdfReadError
import docx
from docx.opc.exceptions import PackageNotFoundError

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB — generous for a text essay


class FileParseError(Exception):
    """Raised for any problem with the uploaded file — always caught and
    turned into a clean 400 response by the /extract-text route, never
    allowed to bubble up as a 500."""
    pass


def _get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[1].lower()


def validate_upload(file_storage) -> str:
    """Checks extension and size before we even try to parse the file.
    Returns the validated extension, or raises FileParseError."""
    if file_storage is None or file_storage.filename == "":
        raise FileParseError("No file was uploaded.")

    ext = _get_extension(file_storage.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise FileParseError(
            f"Unsupported file type '.{ext}'. Please upload a .pdf, .docx, or .txt file."
        )

    # Determine size without permanently consuming the stream
    file_storage.stream.seek(0, io.SEEK_END)
    size = file_storage.stream.tell()
    file_storage.stream.seek(0)
    if size == 0:
        raise FileParseError("The uploaded file is empty.")
    if size > MAX_FILE_SIZE_BYTES:
        raise FileParseError(
            f"File is too large ({size / 1024 / 1024:.1f} MB). Max size is 5 MB."
        )

    return ext


def extract_from_pdf(file_storage) -> tuple[str, list]:
    warnings = []
    try:
        reader = PdfReader(file_storage.stream)
    except PdfReadError:
        raise FileParseError(
            "This PDF could not be read — it may be corrupted or not a valid PDF file."
        )

    if reader.is_encrypted:
        # Try an empty password first (some PDFs are "encrypted" with no
        # real password, just export settings) before giving up.
        try:
            reader.decrypt("")
        except Exception:
            pass
        if reader.is_encrypted:
            raise FileParseError(
                "This PDF is password-protected. Please upload an unlocked version."
            )

    pages_text = []
    image_count = 0
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            # A single malformed page shouldn't take down the whole upload —
            # skip it and keep going with the rest of the document.
            warnings.append("One page could not be read and was skipped.")
        try:
            image_count += len(page.images)
        except Exception:
            pass

    if image_count > 0:
        warnings.append(
            f"{image_count} image(s) in this PDF were skipped — only text is analyzed."
        )

    text = "\n\n".join(p.strip() for p in pages_text if p.strip())
    return text, warnings


def extract_from_docx(file_storage) -> tuple[str, list]:
    warnings = []
    try:
        document = docx.Document(file_storage.stream)
    except PackageNotFoundError:
        raise FileParseError(
            "This .docx file could not be read — it may be corrupted, or "
            "might actually be an old .doc file (only .docx is supported)."
        )
    except Exception:
        raise FileParseError("This .docx file could not be read.")

    # Paragraph text only — this is what naturally "tackles" images without
    # crashing: python-docx's paragraph.text never includes image data,
    # so embedded images are simply invisible to this loop, not a source
    # of errors. We separately count them just to inform the user.
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    try:
        image_count = len(document.inline_shapes)
    except Exception:
        image_count = 0
    if image_count > 0:
        warnings.append(
            f"{image_count} image(s) in this document were skipped — only text is analyzed."
        )

    # Tables aren't walked above (document.paragraphs doesn't include table
    # cell text) — flag this so it's a known limitation, not a silent gap.
    if document.tables:
        warnings.append(
            f"{len(document.tables)} table(s) in this document were not included in the analysis."
        )

    return text, warnings


def extract_from_txt(file_storage) -> tuple[str, list]:
    raw = file_storage.stream.read()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = raw.decode("latin-1")
        except Exception:
            raise FileParseError("Could not read this text file's encoding.")
    return text.strip(), []


def extract_text_from_upload(file_storage) -> dict:
    """
    Main entry point. Validates and extracts text from an uploaded file.
    Returns {"text": str, "warnings": [str, ...]}.
    Raises FileParseError with a user-facing message on any failure —
    the Flask route is expected to catch this and return it as a 400,
    never as an unhandled 500.
    """
    ext = validate_upload(file_storage)

    if ext == "pdf":
        text, warnings = extract_from_pdf(file_storage)
    elif ext == "docx":
        text, warnings = extract_from_docx(file_storage)
    else:  # txt
        text, warnings = extract_from_txt(file_storage)

    if not text or len(text.strip()) == 0:
        raise FileParseError(
            "No readable text was found in this file. If it's a scanned "
            "document (photos of pages), text extraction won't work — "
            "try a document with real text instead."
        )

    return {"text": text, "warnings": warnings}
